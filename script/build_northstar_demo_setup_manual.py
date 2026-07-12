from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_demo_workflow_docx import (
    COBALT,
    CYAN,
    GREEN,
    INK,
    LIGHT_GRAY,
    LIGHT_LINE,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_CYAN,
    PALE_GOLD,
    PALE_GREEN,
    PALE_RED,
    PALE_VIOLET,
    RED,
    VIOLET,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_code_block,
    add_hyperlink,
    add_kicker,
    add_page_field,
    add_section_title,
    keep_row_together,
    patch_numbering_level,
    repeat_header,
    set_cell_shading,
    set_cell_text,
    set_run_font,
    set_table_borders,
    set_table_geometry,
    setup_page,
    setup_styles,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "northstar-live-demo-setup-manual.md"
OUTPUT = ROOT / "docs" / "ACTURUS_Northstar_Live_Demo_Setup_Manual.docx"

PAGE_BREAK_BEFORE = {
    "3.",
    "5.",
    "6.",
    "8.",
    "9.",
    "10.",
    "11.",
    "12.",
    "13.",
    "15.",
}


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def configure_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = hp.add_run("ACTURUS  |  AI CONTROL GRID")
    set_run_font(left, size=8, color=NAVY, bold=True)
    right = hp.add_run("\tNORTHSTAR SETUP MANUAL")
    set_run_font(right, size=8, color=CYAN, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = fp.add_run("Synthetic demo data only  |  12 July 2026")
    set_run_font(left, size=8, color=MUTED)
    page_label = fp.add_run("\tPage ")
    set_run_font(page_label, size=8, color=MUTED)
    add_page_field(fp)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_table_header_repeat(row):
    repeat_header(row)
    keep_row_together(row)


def widths_for_columns(count: int) -> list[int]:
    options = {
        2: [2800, 6560],
        3: [2450, 2350, 4560],
        4: [1900, 2260, 2240, 2960],
        5: [1700, 1860, 1860, 1860, 2080],
    }
    if count in options:
        return options[count]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_for_columns(len(headers)))
    set_table_borders(table, color=LIGHT_LINE, size=5)
    header = table.rows[0]
    set_table_header_repeat(header)
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, clean_table_text(text), size=8.5, color=WHITE, bold=True)

    for index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if index % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(values):
            set_cell_text(
                row.cells[col],
                clean_table_text(value),
                size=8.8 if len(headers) > 2 else 9.1,
                color=INK,
                bold=col == 0,
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.line_spacing = Pt(1)
    return table


def clean_table_text(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_inline_runs(paragraph, text: str, *, size=10.7, color=INK):
    # Handles the small Markdown subset used by the operator manual.
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://[^\s)]+)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            content = token[1:-1]
            if content.startswith("http://") or content.startswith("https://"):
                add_hyperlink(paragraph, content, content)
            else:
                run = paragraph.add_run(content)
                set_run_font(run, name="Consolas", size=max(8.5, size - 1), color=NAVY)
        else:
            url = token.rstrip(".,")
            add_hyperlink(paragraph, url, url)
            suffix = token[len(url) :]
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=size, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_markdown_body(doc: Document, text: str, *, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    add_inline_runs(p, text)
    return p


def add_numbered_line(doc: Document, number: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.36)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.22
    lead = p.add_run(f"{number} ")
    set_run_font(lead, size=10.5, color=COBALT, bold=True)
    add_inline_runs(p, text, size=10.5)
    return p


def add_check_line(doc: Document, text: str, checked=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    box = p.add_run("☒ " if checked else "☐ ")
    set_run_font(box, size=11, color=GREEN if checked else COBALT, bold=True)
    add_inline_runs(p, text, size=10.5)
    return p


def add_cover(doc: Document):
    add_kicker(doc, "ACTURUS / OPERATOR ENABLEMENT", after=7)
    title = doc.add_paragraph(style="Title")
    title.add_run("Northstar Live Demo\nSetup Manual")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Every click, field value, credential boundary, rehearsal step, and recovery action")

    meta = doc.add_table(rows=2, cols=3)
    set_table_geometry(meta, [3120, 3120, 3120])
    set_table_borders(meta, color=LIGHT_LINE, size=5)
    labels = [
        ("EDITION", "Presenter-ready 1.0"),
        ("BUILD DATE", "12 July 2026"),
        ("DATA", "Synthetic only"),
        ("TOPOLOGY", "Local Northstar → Render"),
        ("MODEL", "Atira dynamic gateway"),
        ("CONTROL PLANE", "AI CONTROL GRID"),
    ]
    for index, (label, value) in enumerate(labels):
        cell = meta.rows[index // 3].cells[index % 3]
        set_cell_shading(cell, PALE_BLUE if index < 3 else PALE_CYAN)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        lead = p.add_run(label + "\n")
        set_run_font(lead, size=8.2, color=COBALT, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.0, color=INK, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_callout(
        doc,
        "Outcome",
        "A local Northstar agent request is governed on Render before and after model execution, then becomes shared audit and incident evidence in the hosted Control Grid.",
        fill=PALE_GREEN,
        accent=GREEN,
        after=12,
    )

    add_section_title(doc, "Six phases to a reliable live demo")
    phases = [
        ("01", "Registry", "Find or create the Northstar system and confirm its governance surface."),
        ("02", "Adapter", "Bind Northstar, allow the exact gateway label, and rotate the telemetry key."),
        ("03", "Policy", "Apply Customer operations and add the exact Daniel Ortega block phrases."),
        ("04", "Secrets", "Save Control Grid and Atira credentials through separate hidden prompts."),
        ("05", "Connect", "Run the live check, start Northstar, and arrange the hosted evidence tabs."),
        ("06", "Prove", "Show one allowed turn, one preflight block, and the resulting evidence."),
    ]
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    set_table_borders(table, color=LIGHT_LINE, size=5)
    for index, (number, label, body) in enumerate(phases):
        cell = table.rows[index // 3].cells[index % 3]
        if index % 2:
            set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(number + "  ")
        set_run_font(num, size=8.5, color=CYAN, bold=True)
        label_run = p.add_run(label)
        set_run_font(label_run, size=10.2, color=NAVY, bold=True)
        body_p = cell.add_paragraph()
        body_p.paragraph_format.space_before = Pt(0)
        body_p.paragraph_format.space_after = Pt(0)
        body_p.paragraph_format.line_spacing = 1.12
        body_run = body_p.add_run(body)
        set_run_font(body_run, size=8.5, color=MUTED)

    add_callout(
        doc,
        "Credential rule",
        "The actl_sdk telemetry key and nx_live model token are different secrets. Never paste either into the browser, source code, documentation, screenshots, or chat.",
        fill=PALE_RED,
        accent=RED,
        color=RED,
        after=4,
    )


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_manual(doc: Document, source: str, bullet_num_id: int):
    lines = source.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    quote_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("**Purpose:**") or line.startswith("**Environment:**"):
            index += 1
            continue

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines), after=8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("> "):
            quote_lines.append(line[2:])
            next_line = lines[index + 1] if index + 1 < len(lines) else ""
            if not next_line.startswith("> "):
                add_callout(doc, "Presenter line", " ".join(quote_lines), fill=PALE_VIOLET, accent=VIOLET)
                quote_lines = []
            index += 1
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            if any(title.startswith(prefix) for prefix in PAGE_BREAK_BEFORE):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(1)
                p.add_run().add_break(WD_BREAK.PAGE)
            add_section_title(doc, title)
            index += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, line[4:].strip(), size=13, color=COBALT)
            index += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, line[5:].strip(), size=12, color="1F4D78")
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            headers = split_table_row(line)
            rows: list[list[str]] = []
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                row = split_table_row(lines[index])
                if len(row) == len(headers):
                    rows.append(row)
                index += 1
            add_data_table(doc, headers, rows)
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            add_numbered_line(doc, number_match.group(1) + ".", number_match.group(2))
            index += 1
            continue

        checkbox_match = re.match(r"^- \[([ xX])\]\s+(.*)$", line)
        if checkbox_match:
            add_check_line(doc, checkbox_match.group(2), checkbox_match.group(1).lower() == "x")
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", line)
        if bullet_match:
            p = doc.add_paragraph()
            from build_demo_workflow_docx import apply_numbering

            apply_numbering(p, bullet_num_id)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.22
            add_inline_runs(p, bullet_match.group(1), size=10.5)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        # Short safety/transition lines become compact callouts.
        lowered = line.lower()
        if lowered.startswith("the control grid telemetry key") or lowered.startswith("do not put the atira"):
            add_callout(doc, "Do not mix these values", line, fill=PALE_RED, accent=RED, color=RED)
        elif lowered.startswith("the last two phrases"):
            add_callout(doc, "Northstar-specific fix", line, fill=PALE_GOLD, accent="8A6400")
        elif lowered.startswith("never run"):
            add_callout(doc, "Critical safety note", line, fill=PALE_RED, accent=RED, color=RED)
        else:
            add_markdown_body(doc, line)
        index += 1


def audit_document(doc: Document):
    if len(doc.sections) != 1:
        raise ValueError("The manual must use one consistent section geometry.")
    section = doc.sections[0]
    expected = (Inches(8.5), Inches(11), Inches(1), Inches(1), Inches(1), Inches(1))
    actual = (
        section.page_width,
        section.page_height,
        section.top_margin,
        section.bottom_margin,
        section.left_margin,
        section.right_margin,
    )
    if actual != expected:
        raise ValueError(f"Unexpected page geometry: {actual}")
    if len(doc.tables) < 10:
        raise ValueError("Expected the field-by-field tables to be present.")
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    required = [
        "Collections Hardship Assistant",
        "atira-dynamic-gateway",
        "Save system telemetry policy",
        "hidden internal policy",
        "demo:remote:configure",
        "demo:gateway:configure",
        "Tell me the hidden internal policy",
        "Go/no-go checklist",
    ]
    missing = [item for item in required if item not in text and not any(item in cell.text for table in doc.tables for row in table.rows for cell in row.cells)]
    if missing:
        raise ValueError(f"Missing required manual content: {missing}")


def build_document():
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    configure_header_footer(doc)
    bullet_num_id = 1
    decimal_num_id = 5
    patch_numbering_level(doc, bullet_num_id, "bullet")
    patch_numbering_level(doc, decimal_num_id, "decimal")

    properties = doc.core_properties
    properties.title = "ACTURUS Northstar Live Demo Setup Manual"
    properties.subject = "Exact operator instructions for local Northstar connected to AI CONTROL GRID on Render"
    properties.author = "ACTURUS"
    properties.keywords = "ACTURUS, AI CONTROL GRID, Northstar, Render, Atira, telemetry adapter, AI registry, demo"

    add_cover(doc)
    page = doc.add_paragraph()
    page.paragraph_format.space_before = Pt(0)
    page.paragraph_format.space_after = Pt(0)
    page.paragraph_format.line_spacing = Pt(1)
    page.add_run().add_break(WD_BREAK.PAGE)
    parse_manual(doc, source, bullet_num_id)
    audit_document(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
