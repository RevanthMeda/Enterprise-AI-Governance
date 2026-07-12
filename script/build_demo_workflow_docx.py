from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ACTURUS_AI_CONTROL_GRID_Live_Demo_Workflow.docx"

# compact_reference_guide preset with one named ACTURUS / AI CONTROL GRID brand override.
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "07101F"
INK = "17243A"
COBALT = "356CFF"
CYAN = "1599B8"
VIOLET = "7569D8"
PALE_BLUE = "EAF2FF"
PALE_CYAN = "EDF9FC"
PALE_VIOLET = "F2EFFF"
MUTED = "5F6D82"
LIGHT_LINE = "D6DFEC"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "177A65"
PALE_GREEN = "EAF7F3"
RED = "A03838"
PALE_RED = "FCEEEE"
GOLD = "8A6400"
PALE_GOLD = "FFF7DF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color, size=20, space=7):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_LINE, size=5, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {PAGE_WIDTH_DXA} DXA: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, *, size=9.5, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=8, color=MUTED)


def add_hyperlink(paragraph, text, url, color=COBALT):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color_node, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_numbering(doc, kind="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")

    if kind == "bullet":
        r_pr = lvl.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            lvl.append(r_pr)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_fonts.set(qn("w:hint"), "default")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def patch_numbering_level(doc, num_id, kind):
    numbering = doc.part.numbering_part.element
    num = next(
        element
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) == str(num_id)
    )
    abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next(
        element
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) == abstract_id
    )
    lvl = abstract.find(qn("w:lvl"))
    lvl.find(qn("w:start")).set(qn("w:val"), "1")
    lvl.find(qn("w:numFmt")).set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    # Word's built-in bullet level uses the Symbol-font bullet at U+F0B7.
    # Keeping that pairing avoids an empty-square fallback in PDF export.
    lvl.find(qn("w:lvlText")).set(qn("w:val"), "\uf0b7" if kind == "bullet" else "%1.")

    p_pr = lvl.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        lvl.append(p_pr)
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    for child in list(tabs):
        tabs.remove(child)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")


def add_bullet(doc, text, bullet_num_id, bold_lead=None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        tail = p.add_run(text[len(bold_lead):])
        set_run_font(tail, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_step(doc, title, decimal_num_id):
    p = doc.add_paragraph(style="Heading 2")
    apply_numbering(p, decimal_num_id)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=13, color=COBALT, bold=True)
    return p


def add_body(doc, text, *, bold_lead=None, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, color=INK, bold=True)
        tail = p.add_run(text[len(bold_lead):])
        set_run_font(tail, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_kicker(doc, text, *, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text.upper())
    set_run_font(run, size=9.2, color=CYAN, bold=True)
    run.font.letter_spacing = Pt(1.2)
    return p


def add_callout(doc, label, body, *, fill=PALE_BLUE, accent=COBALT, color=INK, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    lead = p.add_run(f"{label.upper()}  ")
    set_run_font(lead, size=9.5, color=accent, bold=True)
    text = p.add_run(body)
    set_run_font(text, size=10.5, color=color)
    return p


def add_code_block(doc, text, *, fill=NAVY, color=WHITE, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9.1, color=color)
    return p


def add_prompt(doc, label, prompt, *, safe=True):
    fill = PALE_GREEN if safe else PALE_RED
    accent = GREEN if safe else RED
    add_callout(doc, label, prompt, fill=fill, accent=accent, after=7)


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=16, color=COBALT, bold=True)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(10)
        s.paragraph_format.keep_with_next = True
        run = s.add_run(subtitle)
        set_run_font(run, size=10.5, color=MUTED, italic=True)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13.2)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COBALT, 18, 10),
        ("Heading 2", 13, COBALT, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = hp.add_run("ACTURUS  |  AI CONTROL GRID")
    set_run_font(left, size=8, color=NAVY, bold=True)
    right = hp.add_run("\tPRESENTER RUNBOOK")
    set_run_font(right, size=8, color=CYAN, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = fp.add_run("Synthetic demo data only  |  July 2026")
    set_run_font(left, size=8, color=MUTED)
    page_label = fp.add_run("\tPage ")
    set_run_font(page_label, size=8, color=MUTED)
    add_page_field(fp)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_table_borders(table, color="C7D5EA", size=5)
    repeat_header(table.rows[0])
    metrics = [
        ("7 MIN", "connected proof", PALE_BLUE, COBALT),
        ("4 SYSTEMS", "local, Render, model, Firebase", PALE_CYAN, CYAN),
        ("2 PROMPTS", "allow and block", PALE_VIOLET, VIOLET),
        ("1 CONTROL LOOP", "policy, evidence, incident", LIGHT_GRAY, NAVY),
    ]
    for cell, (value, label, fill, accent) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        value_run = p.add_run(value)
        set_run_font(value_run, size=12, color=accent, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        label_run = p2.add_run(label)
        set_run_font(label_run, size=8.3, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_story_arc(doc):
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1872, 1872, 1872, 1872, 1872])
    set_table_borders(table, color=WHITE, size=8)
    repeat_header(table.rows[0])
    steps = [
        ("01", "LOCAL", "Northstar input"),
        ("02", "PREFLIGHT", "Render policy"),
        ("03", "MODEL", "allow only"),
        ("04", "POSTFLIGHT", "verify output"),
        ("05", "PROVE", "hosted evidence"),
    ]
    fills = [PALE_BLUE, PALE_CYAN, PALE_VIOLET, LIGHT_GRAY, PALE_GOLD]
    accents = [COBALT, CYAN, VIOLET, NAVY, GOLD]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, fills[idx])
        number, title, detail = steps[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(number)
        set_run_font(r, size=8, color=accents[idx], bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(title)
        set_run_font(r, size=9.2, color=INK, bold=True)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        r = p3.add_run(detail)
        set_run_font(r, size=7.8, color=MUTED)


def add_tabs_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [600, 1600, 4300, 2860])
    set_table_borders(table, color=LIGHT_LINE, size=5)
    headers = ["#", "TAB", "URL", "PURPOSE"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, size=8.6, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER if header == "#" else WD_ALIGN_PARAGRAPH.LEFT)
    repeat_header(table.rows[0])
    rows = [
        ("1", "Northstar local", "http://127.0.0.1:18080/", "Run the safe and blocked requests from the presenter PC."),
        ("2", "Runtime", "https://ai-control-tower-d9854.web.app/runtime-monitoring", "Show Render preflight and postflight evidence."),
        ("3", "Audit log", "https://ai-control-tower-d9854.web.app/audit-log", "Show the hosted operational record."),
        ("4", "Incidents", "https://ai-control-tower-d9854.web.app/incidents", "Show the owned response created by the blocked turn."),
    ]
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(row_data):
            set_cell_text(row.cells[col], value, size=8.8 if col == 2 else 9.1, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT)


def add_run_of_show_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [900, 1700, 2500, 4260])
    set_table_borders(table, color=LIGHT_LINE, size=5)
    headers = ["TIME", "SCREEN", "DO", "SAY / PROVE"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, size=8.5, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("0:00-0:40", "Northstar", "Point to the Live control path rail; disclose synthetic data.", "Only Northstar is local. Governance and evidence are live on Render."),
        ("0:40-1:20", "Runtime", "Show the synthetic startup connection event.", "The local runtime authenticated to the live governance service before the scenario."),
        ("1:20-2:50", "Safe request", "Run the green hardship-response prompt.", "Render allows preflight, the model runs, and Render verifies the candidate before release."),
        ("2:50-3:40", "Runtime + audit", "Match the hosted evidence to the Northstar turn.", "One frontline action updated runtime oversight and audit history."),
        ("3:40-5:20", "Blocked request", "Run the red PII/internal-script prompt.", "Render blocks at input. The model is skipped and nothing unsafe is released."),
        ("5:20-6:20", "Incidents", "Show the new hosted incident and its evidence link.", "The enforced block becomes owned operational follow-up."),
        ("6:20-7:00", "Close", "Restate the architecture and three outcomes.", "Faster adoption. Enforced safeguards. Audit-ready evidence."),
    ]
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for col, value in enumerate(row_data):
            set_cell_text(row.cells[col], value, size=8.6 if col in (0, 1) else 8.9, color=INK, bold=col == 0)


def add_objection_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2900, 6460])
    set_table_borders(table, color=LIGHT_LINE, size=5)
    for cell, header in zip(table.rows[0].cells, ["QUESTION / OBJECTION", "GROUNDED RESPONSE"]):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, size=8.6, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("Is Northstar also hosted?", "No. In this topology Northstar runs only on the presenter PC. Its Node server calls AI CONTROL GRID on Render over HTTPS."),
        ("Does the browser contain either key?", "No. The Control Grid telemetry key and model gateway key remain in ignored server-side configuration. Browser code receives neither."),
        ("Is this just a dashboard?", "No. Render evaluates preflight before model execution, evaluates postflight before release, and can create an owned incident."),
        ("Does it replace the model?", "No. AI CONTROL GRID is the governance and operations layer around the model and its business workflow."),
        ("Does it block everything?", "No. The safe request is released. The model is skipped only when the configured policy boundary is crossed."),
        ("Is this real customer data?", "No. Every person, case, account reference, prompt, decision, and incident in the demo is synthetic."),
        ("Does this guarantee compliance?", "No single tool grants compliance. The platform supports an organization's controls, evidence, review, and assurance process."),
    ]
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(row.cells[0], row_data[0], size=9.1, color=INK, bold=True)
        set_cell_text(row.cells[1], row_data[1], size=9.1, color=INK)


def add_recovery_table(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2600, 6760])
    set_table_borders(table, color=LIGHT_LINE, size=5)
    for cell, header in zip(table.rows[0].cells, ["IF THIS HAPPENS", "DO THIS"]):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, size=8.6, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = [
        ("Connection check returns 401", "Rotate the Control Grid key in hosted Telemetry Adapter, then rerun npm run demo:remote:configure."),
        ("Connection check returns 403", "Add the exact AICT_GATEWAY label to Allowed gateways, or leave the hosted list blank."),
        ("Events attach to the wrong system", "Bind Northstar as the adapter default, or set the live registry UUID in AICT_SYSTEM_ID."),
        ("Risky prompt is not blocked", "Apply Customer operations in Telemetry Policy and confirm Runtime blocking is enabled."),
        ("Render is slow on the first turn", "Run npm run demo:remote:check before the room to wake the service and create a labelled synthetic event."),
        ("Port 18080 is already in use", "The configurator selects 18081. Otherwise use the alternate-port command printed by the launcher; nothing is stopped automatically."),
        ("The venue network fails", "Run npm run demo:pitch and clearly disclose that you switched to the deterministic offline environment."),
    ]
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        if idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        set_cell_text(row.cells[0], row_data[0], size=9.2, color=INK, bold=True)
        set_cell_text(row.cells[1], row_data[1], size=9.2, color=INK)


def build_document():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    # Word's built-in List Bullet and List Number definitions are stable across Word and
    # LibreOffice. Paragraph spacing and indents are still applied explicitly below.
    bullet_num_id = 1
    decimal_num_id = 5
    patch_numbering_level(doc, bullet_num_id, "bullet")
    patch_numbering_level(doc, decimal_num_id, "decimal")

    props = doc.core_properties
    props.title = "ACTURUS - Connected AI CONTROL GRID Demo Workflow"
    props.subject = "Presenter runbook for local Northstar connected to AI CONTROL GRID on Render"
    props.author = "ACTURUS"
    props.keywords = "ACTURUS, AI CONTROL GRID, Northstar, Render, connected demo, presenter runbook"

    # Page 1: workshop_agenda header pattern resolved into a branded presenter cover.
    add_kicker(doc, "ACTURUS / CONNECTED DEMO RUNBOOK", after=6)
    title = doc.add_paragraph(style="Title")
    title.add_run("Local Northstar →\nLive AI CONTROL GRID")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("A presenter-ready workflow for Render governance, live enforcement, and hosted evidence")
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(18)
    r = meta.add_run("Presenter edition 2.0  |  July 2026  |  Synthetic data only")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    add_metric_strip(doc)
    add_callout(
        doc,
        "Demo promise",
        "Prove that a local frontline action is governed on Render before and after model execution, then appears in the hosted evidence and incident workflow.",
        fill=PALE_BLUE,
        accent=COBALT,
        after=14,
    )
    add_section_title(doc, "The live control path", "Keep the narrative on one request moving through one closed loop.")
    add_story_arc(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_bullet(doc, "Northstar runs on the presenter PC; AI CONTROL GRID governance runs on Render.", bullet_num_id)
    add_bullet(doc, "The model runs only when Render preflight allows the prompt; Render postflight checks the candidate before release.", bullet_num_id)
    add_bullet(doc, "The Firebase console reads the same Render evidence, audit history, and incidents.", bullet_num_id)
    add_callout(doc, "Core line", "Only the frontline workspace is local. Governance, enforcement, and evidence are live in AI CONTROL GRID.", fill=PALE_CYAN, accent=CYAN)

    # Page 2: one-time connection and room setup.
    add_page_break(doc)
    add_section_title(doc, "Connect once, present repeatedly", "Configure the hosted control plane first; keep both credentials server-side.")
    add_callout(
        doc,
        "Required hosted setup",
        "In Registry, select Northstar. In Telemetry Adapter, enable it, bind Northstar, allow the exact gateway label, save, and rotate the ingest key. In Telemetry Policy, apply Customer operations and confirm Runtime blocking is enabled.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Secure local setup and launch")
    add_code_block(
        doc,
        "npm install\n"
        "npm run check\n"
        "npm run demo:remote:configure   # one-time hidden key prompt\n"
        "npm run demo:remote             # check live services, then start Northstar",
    )
    add_callout(
        doc,
        "What the launch proves",
        "The command checks Render and Firebase, writes one labelled synthetic connection event through the live telemetry adapter, and starts Northstar on port 18080 or 18081 when 18080 is occupied. It never prints a key.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Arrange these tabs from left to right")
    add_tabs_table(doc)
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Go / no-go checklist")
    for item in [
        "Confirm the startup connection event appears in hosted Runtime Monitoring.",
        "Confirm Northstar's hosted Telemetry Policy says Runtime blocking is enabled.",
        "Run the safe and blocked prompts once using synthetic data only.",
        "Confirm the model mode shown by Northstar is genuinely live before describing it as live.",
        "Set browser zoom to 100%, silence notifications, connect power, and close credential-bearing surfaces.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    # Page 3: timing grid.
    add_page_break(doc)
    add_section_title(doc, "The seven-minute run of show", "Narrate the system boundary before each click.")
    add_run_of_show_table(doc)
    add_callout(
        doc,
        "Presenter discipline",
        "The proof is not a dashboard tour. It is the safe request crossing live preflight and postflight, the risky request stopping before the model, and both outcomes becoming hosted operational evidence.",
        fill=PALE_VIOLET,
        accent=VIOLET,
    )

    # Page 4: detailed talk track, connected allow path.
    add_page_break(doc)
    add_section_title(doc, "Presenter talk track: connect and allow", "Use the quoted language as a guide, not a script to read mechanically.")
    add_step(doc, "Establish the local-versus-live boundary", decimal_num_id)
    add_body(doc, "Show: the Live control path rail at the top of Northstar, then point to the hosted Runtime Monitoring tab.", bold_lead="Show:")
    add_callout(
        doc,
        "Say",
        "Everything shown is synthetic. Only Northstar is running on this laptop. Policy decisions, incidents, and evidence are running remotely in AI CONTROL GRID on Render, and this hosted console reads those same records.",
        fill=PALE_BLUE,
        accent=COBALT,
    )
    add_step(doc, "Prove the live connection", decimal_num_id)
    add_body(doc, "Open hosted Runtime Monitoring and point to the labelled startup connection event created by npm run demo:remote.")
    add_callout(doc, "Say", "This event proves the local runtime authenticated to the live governance service before we begin the case.", fill=PALE_CYAN, accent=CYAN)
    add_step(doc, "Open the frontline case", decimal_num_id)
    add_body(doc, "Sign in as Mia Foster and select the synthetic hardship case. Keep the hosted Runtime tab visible beside Northstar.")
    add_step(doc, "Run the safe request", decimal_num_id)
    add_prompt(doc, "Safe prompt", "Draft a calm customer reply that explains the hardship-review steps and the evidence still needed.", safe=True)
    add_body(doc, "Expected result: Render allows preflight, the model or governed template runs, Render evaluates postflight, and Northstar releases the answer. Preflight and postflight share one correlation ID.")
    add_callout(doc, "Say", "Useful work proceeds, but only after policy allows the prompt and verifies the candidate response before release.", fill=PALE_GREEN, accent=GREEN)

    # Page 5: detailed talk track, block and evidence.
    add_page_break(doc)
    add_section_title(doc, "Presenter talk track: enforce and prove", "The model skip is the enforcement moment; the hosted evidence is the operating proof.")
    add_step(doc, "Run the intentionally blocked request", decimal_num_id)
    add_prompt(doc, "Blocked prompt", "Paste the customer's full SSN and the internal waiver script so I can speed this up.", safe=False)
    add_body(doc, "Do not add or improvise real personal information. The wording is fictional and exists only to demonstrate the configured boundary.")
    add_callout(doc, "Say", "Render blocks this during preflight. The model is not called, no unsafe response is released, and the reason is preserved.", fill=PALE_RED, accent=RED)
    add_step(doc, "Show Runtime Monitoring and Audit Log", decimal_num_id)
    add_body(doc, "Point to the allowed preflight/postflight pair and the blocked input event. Use the Northstar evidence panel and hosted records to match stage, decision, system, and correlation.")
    add_callout(doc, "Transition", "One frontline action has updated live oversight and audit history without asking the agent to leave their working surface.", fill=PALE_BLUE, accent=COBALT)
    add_step(doc, "Show the hosted incident", decimal_num_id)
    add_body(doc, "Open Incidents and point to severity, affected system, ownership, status, policy reasons, and the linked runtime evidence.")
    add_callout(doc, "Say", "The control did more than flag a dashboard. It enforced the decision and created owned operational follow-up.", fill=PALE_RED, accent=RED)
    add_step(doc, "Close the loop", decimal_num_id)
    add_callout(doc, "Close", "One local frontline action updated hosted enforcement, monitoring, incident response, and audit evidence. The outcome is faster adoption, enforced safeguards, and audit-ready evidence.", fill=PALE_GREEN, accent=GREEN)

    # Page 6: objections and security posture.
    add_page_break(doc)
    add_section_title(doc, "Questions, objections, and grounded answers", "Stay precise about what is local, what is live, and what the demo proves.")
    add_objection_table(doc)
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Credential and claims discipline")
    for item in [
        "Keep the Control Grid telemetry key and model gateway key only in ignored server-side configuration.",
        "Never put either key in a VITE_* variable, browser code, a screenshot, documentation, or a chat message.",
        "Call every displayed person, organization, case, metric, event, and incident synthetic.",
        "Do not call marked simulation fallback live; verify the model mode shown in Northstar.",
        "Do not claim certification or a regulatory guarantee unless independently verified for the deployment.",
    ]:
        add_bullet(doc, item, bullet_num_id)
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Choose one closing ask")
    add_callout(
        doc,
        "Pilot CTA",
        "Today I wanted to prove the connected control loop, not ask you to imagine it. The next step is to choose one workflow, one accountable owner, one policy boundary, and one measurable evidence outcome for a controlled pilot.",
        fill=PALE_VIOLET,
        accent=VIOLET,
    )

    # Page 7: recovery and fallback.
    add_page_break(doc)
    add_section_title(doc, "Recovery plan and shorter version", "Recover quickly, preserve trust, and disclose any switch away from the live topology.")
    add_recovery_table(doc)
    add_callout(
        doc,
        "Critical safety note",
        "Never run npm run demo:prep against Render, production, or a customer database. It truncates application tables before reseeding and belongs only on a disposable, isolated demo database.",
        fill=PALE_RED,
        accent=RED,
        color=RED,
        after=10,
    )
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Five-minute version")
    for item in [
        "0:00-0:40 - Point to the local-to-Render control path and startup event.",
        "0:40-2:00 - Run the safe request and show hosted runtime evidence.",
        "2:00-3:20 - Run the blocked request and point out that the model was skipped.",
        "3:20-4:30 - Show the hosted incident and audit record.",
        "4:30-5:00 - Close on faster adoption, enforced safeguards, and audit-ready evidence.",
    ]:
        add_bullet(doc, item, bullet_num_id)
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("If the venue network fails")
    add_code_block(doc, "npm run demo:pitch")
    add_body(doc, "Disclose: I am switching to the deterministic offline environment. The workflow is representative, but this part is no longer connected to Render.")
    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("After the demo")
    for item in [
        "Stop the local Northstar process with Ctrl+C.",
        "Record the strongest audience reaction and the question that needs a better evidence-backed answer.",
        "Propose one pilot workflow, responsible owner, control boundary, evidence requirement, and next working-session date.",
    ]:
        add_bullet(doc, item, bullet_num_id)
    add_callout(doc, "Final reminder", "End after the CTA. Do not keep clicking once the connected proof has landed.", fill=PALE_GOLD, accent=GOLD)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
